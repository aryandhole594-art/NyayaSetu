import io
import os
import re
import tempfile
from dataclasses import dataclass
from datetime import datetime
from typing import Any

import fitz  # PyMuPDF
import pytesseract
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from pytesseract import TesseractNotFoundError


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
DATE_PATTERNS = [
    r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
    r"\b\d{1,2}\s+(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{2,4}\b",
    r"\b(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{1,2},\s*\d{2,4}\b",
]


RISKY_CLAUSE_PATTERNS = [
    {"id": "non_compete_over_1y", "label": "Non-compete exceeds 1 year", "severity": "high", "pattern": r"non[-\s]?compete.{0,120}(?:1[3-9]|[2-9]\d)\s*(?:month|months|year|years)", "query": "freedom of occupation Article 19(1)(g) reasonable restriction"},
    {"id": "non_compete_no_geo_limit", "label": "Non-compete without geographic limit", "severity": "high", "pattern": r"non[-\s]?compete.{0,140}(worldwide|global|anywhere|without geographic limitation)", "query": "restraint of trade reasonableness"},
    {"id": "non_solicit_overbroad", "label": "Overbroad non-solicitation", "severity": "medium", "pattern": r"non[-\s]?solicit.{0,160}(all customers|any customer|potential customers)", "query": "proportional restriction employment contracts"},
    {"id": "uni_termination_no_notice", "label": "Unilateral termination without notice", "severity": "high", "pattern": r"(?:company|employer|licensor).{0,80}terminate.{0,100}(without notice|immediately)", "query": "natural justice termination notice"},
    {"id": "termination_for_convenience_one_sided", "label": "One-sided termination for convenience", "severity": "high", "pattern": r"(?:company|service provider).{0,60}terminate for convenience", "query": "arbitrariness Article 14 contract"},
    {"id": "liquidated_damages_unilateral", "label": "Unilateral liquidated damages", "severity": "medium", "pattern": r"liquidated damages.{0,120}(sole discretion|as determined by company)", "query": "penalty clause fairness"},
    {"id": "ip_assignment_without_comp", "label": "IP assignment without compensation", "severity": "high", "pattern": r"assign(?:ment)? of (?:all )?intellectual property.{0,160}(without compensation|no additional compensation|royalty[-\s]?free)", "query": "fair remuneration for assignment of work"},
    {"id": "moral_rights_waiver", "label": "Waiver of moral rights", "severity": "medium", "pattern": r"waive.{0,80}moral rights", "query": "author rights waiver"},
    {"id": "perpetual_license_irrevocable", "label": "Perpetual irrevocable license", "severity": "medium", "pattern": r"perpetual.{0,80}irrevocable.{0,80}license", "query": "unconscionable IP license"},
    {"id": "mandatory_arbitration_other_state", "label": "Mandatory arbitration in different state", "severity": "high", "pattern": r"arbitration.{0,160}(seat|venue).{0,80}(?:outside|other than).{0,80}(state|city)", "query": "access to justice forum selection unfair"},
    {"id": "exclusive_jurisdiction_remote", "label": "Exclusive jurisdiction in remote court", "severity": "high", "pattern": r"exclusive jurisdiction.{0,120}(only|solely).{0,120}(courts? at)", "query": "forum selection clause fairness"},
    {"id": "waiver_class_action", "label": "Class action waiver", "severity": "medium", "pattern": r"waive.{0,120}(class action|representative action)", "query": "consumer collective redress"},
    {"id": "auto_renewal_silent", "label": "Silent auto-renewal", "severity": "medium", "pattern": r"automatically renew.{0,120}(unless terminated).{0,120}(notice)", "query": "consumer notice and auto renewal"},
    {"id": "price_change_unilateral", "label": "Unilateral price change right", "severity": "high", "pattern": r"(?:may|can) change (?:fees|pricing|charges).{0,120}(at any time|without notice)", "query": "unfair contract terms unilateral variation"},
    {"id": "service_change_unilateral", "label": "Unilateral scope/service change", "severity": "medium", "pattern": r"modify.{0,80}(services|scope).{0,120}(at any time|sole discretion)", "query": "unilateral modification fairness"},
    {"id": "liability_cap_too_low", "label": "Liability cap too low", "severity": "medium", "pattern": r"liability.{0,120}(shall not exceed|limited to).{0,80}(fees paid in last (?:1|one|3|three) month)", "query": "reasonable liability cap"},
    {"id": "exclude_gross_negligence", "label": "Excludes liability for gross negligence", "severity": "high", "pattern": r"no liability.{0,140}(gross negligence|willful misconduct)", "query": "public policy exclusion of gross negligence"},
    {"id": "broad_indemnity_one_sided", "label": "One-sided broad indemnity", "severity": "high", "pattern": r"indemnify.{0,160}(any and all claims|all losses).{0,80}(including attorney|legal fees)", "query": "one sided indemnity fairness"},
    {"id": "indemnity_no_cap", "label": "Indemnity without cap", "severity": "medium", "pattern": r"indemnity.{0,120}(unlimited|without limit)", "query": "proportional indemnity obligations"},
    {"id": "confidentiality_perpetual", "label": "Perpetual confidentiality", "severity": "low", "pattern": r"confidentiality.{0,120}(perpetual|in perpetuity)", "query": "reasonable confidentiality duration"},
    {"id": "audit_intrusive", "label": "Intrusive audit rights", "severity": "medium", "pattern": r"audit.{0,140}(any time|without notice|all records)", "query": "privacy and proportional audit rights"},
    {"id": "assignment_without_consent", "label": "Assignment without consent", "severity": "medium", "pattern": r"(?:company|provider).{0,80}assign.{0,80}without consent", "query": "assignment clause fairness"},
    {"id": "no_injunctive_relief", "label": "No injunctive relief permitted", "severity": "high", "pattern": r"no (?:right to )?injunctive relief", "query": "effective remedy access to courts"},
    {"id": "short_claim_window", "label": "Very short claim window", "severity": "medium", "pattern": r"claim.{0,80}within\s+(?:7|10|15|30)\s+days", "query": "limitation period reasonableness"},
    {"id": "penalty_for_termination_by_user", "label": "Penalty for user termination", "severity": "medium", "pattern": r"termination.{0,120}(penalty|early termination fee)", "query": "penalty and unfair contract terms"},
    {"id": "non_disparagement_perpetual", "label": "Perpetual non-disparagement", "severity": "low", "pattern": r"non[-\s]?disparagement.{0,120}(perpetual|in perpetuity)", "query": "free speech contractual limitation"},
    {"id": "forced_waiver_statutory_rights", "label": "Waiver of statutory rights", "severity": "high", "pattern": r"waive.{0,120}(statutory rights|rights under applicable law)", "query": "cannot waive statutory protections"},
    {"id": "surveillance_consent_broad", "label": "Overbroad monitoring consent", "severity": "medium", "pattern": r"monitor.{0,160}(all communications|all activity|without notice)", "query": "privacy proportionality and consent"},
    {"id": "ownership_of_user_data", "label": "Provider claims ownership of user data", "severity": "high", "pattern": r"(?:own|ownership).{0,120}(user data|customer data)", "query": "data protection informational privacy"},
    {"id": "payment_withhold_unilateral", "label": "Unilateral payment withholding", "severity": "medium", "pattern": r"withhold (?:payment|fees).{0,120}(sole discretion|without notice)", "query": "arbitrary withholding fairness"},
]


@dataclass
class UploadedDocument:
    doc_id: str
    name: str
    text: str
    added_at: str


class DocumentCorpusStore:
    def __init__(self):
        self._docs: dict[str, UploadedDocument] = {}

    def add(self, name: str, text: str) -> UploadedDocument:
        doc_id = f"doc_{len(self._docs) + 1}_{int(datetime.utcnow().timestamp())}"
        doc = UploadedDocument(doc_id=doc_id, name=name, text=text, added_at=datetime.utcnow().isoformat())
        self._docs[doc_id] = doc
        return doc

    def list_docs(self) -> list[dict[str, str]]:
        return [
            {"doc_id": d.doc_id, "name": d.name, "chars": str(len(d.text)), "added_at": d.added_at}
            for d in self._docs.values()
        ]

    def combined_text(self) -> str:
        parts = []
        for d in self._docs.values():
            parts.append(f"\n\n[UPLOADED DOCUMENT: {d.name}]\n{d.text}\n")
        return "\n".join(parts)


def _pdf_page_to_text(page: fitz.Page) -> str:
    return (page.get_text("text") or "").strip()


def _contains_meaningful_text(text: str, min_chars: int = 45) -> bool:
    cleaned = re.sub(r"\s+", "", text or "")
    return len(cleaned) >= min_chars


def _preprocess_for_ocr(img: Image.Image) -> Image.Image:
    # Gentle denoise + contrast + threshold, useful for scans and handwriting.
    gray = ImageOps.grayscale(img)
    denoised = gray.filter(ImageFilter.MedianFilter(size=3))
    boosted = ImageEnhance.Contrast(denoised).enhance(2.0)
    sharpened = boosted.filter(ImageFilter.SHARPEN)
    return sharpened.point(lambda px: 255 if px > 165 else 0)


def _run_ocr(img: Image.Image) -> str:
    configs = [
        "--oem 3 --psm 6",
        "--oem 3 --psm 11",
        "--oem 1 --psm 4",
    ]
    candidates: list[str] = []
    for cfg in configs:
        try:
            text = (pytesseract.image_to_string(img, config=cfg) or "").strip()
            if text:
                candidates.append(text)
        except TesseractNotFoundError:
            raise
        except Exception:
            continue
    if not candidates:
        return ""
    return max(candidates, key=lambda t: len(re.sub(r"\s+", "", t)))


def _pdf_page_to_ocr_text(page: fitz.Page) -> str:
    # Higher raster DPI improves OCR quality for scanned docs.
    pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    raw = _run_ocr(img)
    enhanced = _run_ocr(_preprocess_for_ocr(img))
    return max([raw, enhanced], key=lambda t: len(re.sub(r"\s+", "", t or ""))).strip()


def _extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts: list[str] = []

    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = re.sub(r"\s+", " ", cell.text or "").strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_text(file: str | os.PathLike | Any) -> str:
    """
    Extract text from:
    - PDF (PyMuPDF native extraction)
    - Scanned/handwritten PDF (Tesseract OCR fallback per page)
    - Image files (Tesseract OCR)
    - DOCX (paragraph + table text)
    """
    file_path: str
    temp_path: str | None = None

    if hasattr(file, "read") and not isinstance(file, (str, os.PathLike)):
        suffix = os.path.splitext(getattr(file, "name", "upload.bin"))[1] or ".bin"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file.read())
            temp_path = tmp.name
        file_path = temp_path
    else:
        file_path = os.fspath(file)

    try:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            doc = fitz.open(file_path)
            out: list[str] = []
            for idx, page in enumerate(doc):
                native = _pdf_page_to_text(page)
                native_ok = _contains_meaningful_text(native)
                ocr_text = ""

                if not native_ok:
                    ocr_text = _pdf_page_to_ocr_text(page)
                elif len(re.sub(r"\s+", "", native)) < 220:
                    # For low-text pages, OCR can still recover missed pieces.
                    ocr_text = _pdf_page_to_ocr_text(page)

                chosen = native if native_ok else ""
                if ocr_text and len(re.sub(r"\s+", "", ocr_text)) > len(re.sub(r"\s+", "", chosen)):
                    chosen = ocr_text
                if chosen:
                    out.append(chosen)
                else:
                    out.append(f"[Page {idx + 1}] No readable text extracted.")
            return "\n\n".join(out).strip()

        if ext in IMAGE_EXTENSIONS:
            img = Image.open(file_path)
            raw = _run_ocr(img)
            enhanced = _run_ocr(_preprocess_for_ocr(img))
            best = max([raw, enhanced], key=lambda t: len(re.sub(r"\s+", "", t or "")))
            return best.strip()

        if ext == ".docx":
            return _extract_text_from_docx(file_path)

        raise ValueError("Unsupported file type. Use PDF, DOCX, or image formats.")
    except TesseractNotFoundError as e:
        raise RuntimeError(
            "Tesseract OCR binary not found. Install `tesseract` on the host system to read scanned PDFs/images."
        ) from e
    finally:
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)


def analyze_contract(text: str, rag_index, top_k: int = 2) -> dict[str, Any]:
    findings = []
    for item in RISKY_CLAUSE_PATTERNS:
        match = re.search(item["pattern"], text, flags=re.IGNORECASE | re.DOTALL)
        if not match:
            continue

        start = max(0, match.start() - 120)
        end = min(len(text), match.end() + 120)
        snippet = re.sub(r"\s+", " ", text[start:end]).strip()

        contexts = []
        try:
            contexts = rag_index.retrieve(item["query"], top_k=top_k)
        except Exception:
            contexts = []

        findings.append({
            "id": item["id"],
            "clause": item["label"],
            "severity": item["severity"],
            "matched_text": match.group(0)[:220],
            "snippet": snippet,
            "constitutional_grounding": [
                {
                    "title": c["title"],
                    "article_numbers": c.get("article_numbers", []),
                    "excerpt": c["text"][:240],
                }
                for c in contexts
            ],
        })

    findings.sort(key=lambda f: {"high": 0, "medium": 1, "low": 2}.get(f["severity"], 3))
    return {
        "total_patterns_checked": len(RISKY_CLAUSE_PATTERNS),
        "total_flags": len(findings),
        "flags": findings,
        "explainability": "Pattern match + constitutional RAG grounding (no black-box ML).",
    }


def extract_deadlines(text: str) -> list[dict[str, str]]:
    hits = []
    lower = text.lower()

    for patt in DATE_PATTERNS:
        for m in re.finditer(patt, text, flags=re.IGNORECASE):
            start = max(0, m.start() - 80)
            end = min(len(text), m.end() + 80)
            window = re.sub(r"\s+", " ", text[start:end]).strip()
            urgency = "deadline" if any(k in window.lower() for k in ["before", "within", "last date", "deadline"]) else "date_mentioned"
            hits.append({"date": m.group(0), "type": urgency, "context": window})

    for m in re.finditer(r"within\s+(\d{1,3})\s+(day|days|week|weeks|month|months)", lower):
        start = max(0, m.start() - 60)
        end = min(len(text), m.end() + 60)
        hits.append({
            "date": m.group(0),
            "type": "relative_deadline",
            "context": re.sub(r"\s+", " ", text[start:end]).strip(),
        })

    seen = set()
    deduped = []
    for h in hits:
        key = (h["date"].lower(), h["context"][:80].lower())
        if key in seen:
            continue
        seen.add(key)
        deduped.append(h)
    return deduped[:20]


def extract_legal_sections(text: str) -> list[str]:
    refs = set()
    for m in re.finditer(r"\b(Article\s+\d+[A-Z]?|Section\s+\d+[A-Z]?)\b", text, flags=re.IGNORECASE):
        refs.add(m.group(0).strip())
    return sorted(refs)


def decode_court_notice(text: str, rag_index, top_k: int = 2) -> dict[str, Any]:
    deadlines = extract_deadlines(text)
    sections = extract_legal_sections(text)

    section_explanations = []
    for sec in sections:
        q = f"Explain {sec} in plain language and what action is expected"
        try:
            chunks = rag_index.retrieve(q, top_k=top_k)
        except Exception:
            chunks = []

        section_explanations.append({
            "section": sec,
            "grounding": [
                {
                    "title": c["title"],
                    "article_numbers": c.get("article_numbers", []),
                    "excerpt": c["text"][:260],
                }
                for c in chunks
            ],
        })

    actions = []
    if deadlines:
        actions.append("Track every extracted date and file your reply before the earliest deadline.")
    else:
        actions.append("No explicit deadline found. Verify timeline from the issuing court immediately.")
    if sections:
        actions.append("Read the cited legal sections and align your response to those requirements.")
    actions.append("Prepare supporting documents and consult a qualified lawyer for court-specific procedure.")

    return {
        "deadlines": deadlines,
        "sections_cited": sections,
        "section_explanations": section_explanations,
        "plain_english": {
            "what_to_do": actions,
            "by_when": deadlines[0]["date"] if deadlines else "As soon as possible",
        },
    }
