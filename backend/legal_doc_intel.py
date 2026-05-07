import json
import os
import re
import tempfile
from collections import defaultdict
from datetime import datetime
from dataclasses import dataclass
from typing import Any, Callable

import fitz
import pytesseract
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from pytesseract import TesseractNotFoundError

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except Exception:
    PDFPLUMBER_AVAILABLE = False
    pdfplumber = None

try:
    from langdetect import DetectorFactory, LangDetectException, detect

    LANGDETECT_AVAILABLE = True
    DetectorFactory.seed = 0
except Exception:
    LANGDETECT_AVAILABLE = False
    DetectorFactory = None
    LangDetectException = Exception
    detect = None

try:
    from pdf2image import convert_from_path

    PDF2IMAGE_AVAILABLE = True
except Exception:
    PDF2IMAGE_AVAILABLE = False
    convert_from_path = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False
    A4 = None
    canvas = None

try:
    import cv2
    import numpy as np

    CV2_AVAILABLE = True
except Exception:
    CV2_AVAILABLE = False
    cv2 = None
    np = None

try:
    import magic

    MAGIC_AVAILABLE = True
except Exception:
    MAGIC_AVAILABLE = False
    magic = None

try:
    from document_intelligence import analyze_contract
except Exception:
    analyze_contract = None


SYSTEM_PROMPT = """You are a senior legal document analyst with expertise in contracts, court filings, compliance documents, NDAs, agreements, and regulatory paperwork.

Your job is to:
1. Extract all text content from the provided document
2. Clean and structure the raw extracted text
3. Produce a precise, legally-accurate summary

Return strict JSON only in the format:
{
  "document_overview": {
    "document_type": "",
    "title": "",
    "date": "",
    "jurisdiction": "",
    "total_pages": ""
  },
  "parties_involved": [],
  "purpose_scope": "",
  "key_terms_definitions": [],
  "critical_clauses": [],
  "red_flags_risk_analysis": [],
  "important_dates_deadlines": [],
  "financial_obligations": [],
  "obligations_summary": {
    "party_a_must": [],
    "party_b_must": []
  },
  "attachments_exhibits": [],
  "plain_english_summary": "",
  "disclaimer": "This analysis is AI-generated and for informational purposes only. It does not constitute legal advice. Consult a qualified attorney for legal decisions."
}

Rules:
- Preserve clause numbers and section headers as written.
- Extract dates, party names, defined terms, and monetary values verbatim.
- If unclear, state unclear. Never invent facts.
"""

RED_FLAG_TERMS = [
    "INDEMNIF",
    "ARBITRAT",
    "LIQUIDAT",
    "WAIVER",
    "IRREVOCABLE",
]

DATE_REGEX = re.compile(
    r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|"
    r"\d{1,2}\s+(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{2,4}|"
    r"(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{1,2},\s*\d{2,4})\b",
    flags=re.IGNORECASE,
)

MONEY_REGEX = re.compile(
    r"(?:₹\s?[\d,]+(?:\.\d{1,2})?|INR\s?[\d,]+(?:\.\d{1,2})?|USD\s?\$?[\d,]+(?:\.\d{1,2})?|\$[\d,]+(?:\.\d{1,2})?)"
)


@dataclass
class ExtractionResult:
    file_type: str
    total_pages: int
    raw_text: str
    cleaned_text: str
    extraction_issues: list[str]
    low_quality_scan: bool
    ocr_confidence: float | None
    language: str
    word_count: int
    red_flag_terms: list[str]
    chunks: list[str]
    segments: list[dict[str, Any]]


def _detect_file_type(file_path: str, filename: str = "") -> str:
    ext = os.path.splitext((filename or file_path).lower())[1]
    if ext in {".pdf"}:
        return "pdf"
    if ext in {".png", ".jpg", ".jpeg"}:
        return "image"
    if ext == ".docx":
        return "docx"

    if MAGIC_AVAILABLE:
        try:
            mime = magic.from_file(file_path, mime=True) or ""
            if "pdf" in mime:
                return "pdf"
            if "wordprocessingml.document" in mime:
                return "docx"
            if "image" in mime:
                return "image"
        except Exception:
            pass
    return "unknown"


def _safe_detect_lang(text: str) -> str:
    if not LANGDETECT_AVAILABLE:
        return "unknown"
    if len(text.strip()) < 40:
        return "unknown"
    try:
        return detect(text)
    except LangDetectException:
        return "unknown"


def _remove_headers_footers(text: str) -> str:
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        s = line.strip()
        if not s:
            cleaned.append("")
            continue
        if re.fullmatch(r"page\s+\d+(\s+of\s+\d+)?", s, flags=re.IGNORECASE):
            continue
        if re.fullmatch(r"\d+", s):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)


def _clean_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = _remove_headers_footers(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_chunks_preserve_paragraphs(text: str, chunk_words: int = 2200, overlap_words: int = 150) -> list[str]:
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_words = 0

    for para in paragraphs:
        pw = len(para.split())
        if current and current_words + pw > chunk_words:
            chunks.append("\n\n".join(current))
            overlap = " ".join(" ".join(current).split()[-overlap_words:])
            current = [overlap, para] if overlap else [para]
            current_words = len(" ".join(current).split())
        else:
            current.append(para)
            current_words += pw

    if current:
        chunks.append("\n\n".join(current))
    return chunks or [text]


def _preprocess_image(img: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(img)
    denoised = gray.filter(ImageFilter.MedianFilter(size=3))
    boosted = ImageEnhance.Contrast(denoised).enhance(2.0)
    if CV2_AVAILABLE:
        arr = np.array(boosted)
        arr = cv2.GaussianBlur(arr, (3, 3), 0)
        _, arr = cv2.threshold(arr, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        return Image.fromarray(arr)
    return boosted.point(lambda p: 255 if p > 170 else 0)


def _ocr_with_confidence(img: Image.Image) -> tuple[str, float | None]:
    try:
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT, config="--psm 6")
        text = pytesseract.image_to_string(img, config="--psm 6")
        conf_values = []
        for c in data.get("conf", []):
            try:
                v = float(c)
                if v >= 0:
                    conf_values.append(v)
            except Exception:
                continue
        conf = round(sum(conf_values) / len(conf_values), 2) if conf_values else None
        return text.strip(), conf
    except TesseractNotFoundError:
        raise
    except Exception:
        return "", None


def _extract_pdf(pdf_path: str) -> tuple[str, int, list[str], float | None, list[dict[str, Any]]]:
    issues: list[str] = []
    page_texts: list[str] = []
    ocr_conf_samples: list[float] = []
    segments: list[dict[str, Any]] = []

    with fitz.open(pdf_path) as fdoc:
        page_count = len(fdoc)

    if PDFPLUMBER_AVAILABLE:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                direct = (page.extract_text() or "").strip()
                if len(re.sub(r"\s+", "", direct)) >= 40:
                    page_texts.append(direct)
                    segments.append({"page": i + 1, "text": direct, "source": "pdfplumber"})
                    continue

                if not PDF2IMAGE_AVAILABLE:
                    issues.append(f"[EXTRACTION_ISSUE] Page {i + 1}: scanned/OCR fallback unavailable (pdf2image missing).")
                    page_texts.append(direct)
                    if direct:
                        segments.append({"page": i + 1, "text": direct, "source": "pdfplumber"})
                    continue

                try:
                    images = convert_from_path(pdf_path, dpi=300, first_page=i + 1, last_page=i + 1)
                    if not images:
                        issues.append(f"[EXTRACTION_ISSUE] Page {i + 1}: could not rasterize for OCR.")
                        page_texts.append("")
                        continue
                    prepped = _preprocess_image(images[0])
                    ocr_text, conf = _ocr_with_confidence(prepped)
                    if conf is not None:
                        ocr_conf_samples.append(conf)
                    if not ocr_text:
                        issues.append(f"[EXTRACTION_ISSUE] Page {i + 1}: OCR returned no readable text.")
                    page_texts.append(ocr_text)
                    segments.append({"page": i + 1, "text": ocr_text, "source": "ocr"})
                except TesseractNotFoundError:
                    raise
                except Exception as e:
                    issues.append(f"[EXTRACTION_ISSUE] Page {i + 1}: OCR failed ({e}).")
                    page_texts.append("")
                    segments.append({"page": i + 1, "text": "", "source": "ocr"})
    else:
        issues.append("[EXTRACTION_ISSUE] pdfplumber not available; using PyMuPDF fallback extraction.")
        with fitz.open(pdf_path) as doc:
            for i, page in enumerate(doc):
                direct = (page.get_text("text") or "").strip()
                if len(re.sub(r"\s+", "", direct)) >= 40:
                    page_texts.append(direct)
                    segments.append({"page": i + 1, "text": direct, "source": "pymupdf"})
                    continue
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                prepped = _preprocess_image(img)
                ocr_text, conf = _ocr_with_confidence(prepped)
                if conf is not None:
                    ocr_conf_samples.append(conf)
                if not ocr_text:
                    issues.append(f"[EXTRACTION_ISSUE] Page {i + 1}: OCR returned no readable text.")
                page_texts.append(ocr_text)
                segments.append({"page": i + 1, "text": ocr_text, "source": "ocr"})

    avg_conf = round(sum(ocr_conf_samples) / len(ocr_conf_samples), 2) if ocr_conf_samples else None
    return "\n\n".join(page_texts).strip(), page_count, issues, avg_conf, segments


def _extract_image(image_path: str) -> tuple[str, list[str], float | None, list[dict[str, Any]]]:
    issues: list[str] = []
    img = Image.open(image_path)
    text_raw, conf_raw = _ocr_with_confidence(img)
    text_pre, conf_pre = _ocr_with_confidence(_preprocess_image(img))

    if len(re.sub(r"\s+", "", text_pre)) > len(re.sub(r"\s+", "", text_raw)):
        text = text_pre
        conf = conf_pre
    else:
        text = text_raw
        conf = conf_raw

    if not text:
        issues.append("[EXTRACTION_ISSUE] OCR returned no readable text from image.")
    return text, issues, conf, [{"page": 1, "text": text, "source": "ocr"}]


def _extract_docx(docx_path: str) -> tuple[str, int, list[str], list[dict[str, Any]]]:
    issues: list[str] = []
    doc = Document(docx_path)
    lines: list[str] = []
    segments: list[dict[str, Any]] = []

    for section in doc.sections:
        header_text = " ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
        if header_text:
            lines.append(f"[HEADER] {header_text}")
            segments.append({"page": "DOCX", "text": header_text, "source": "header"})

    for idx, p in enumerate(doc.paragraphs, start=1):
        txt = (p.text or "").strip()
        if not txt:
            continue
        style = (p.style.name or "").strip()
        if style.lower().startswith("heading"):
            lines.append(f"[{style}] {txt}")
        else:
            lines.append(txt)
        segments.append({"page": "DOCX", "paragraph": idx, "text": txt, "source": style or "paragraph"})

    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            row_cells = []
            raw_cells = []
            for cell in row.cells:
                cell_text = re.sub(r"\s+", " ", cell.text or "").strip()
                if cell_text:
                    row_cells.append(cell_text)
                    raw_cells.append(cell_text)
            if row_cells:
                lines.append(" | ".join(row_cells))
                segments.append(
                    {
                        "page": "DOCX",
                        "table": table_index,
                        "row": row_index,
                        "text": " | ".join(raw_cells),
                        "source": "table",
                    }
                )

    if not lines:
        issues.append("[EXTRACTION_ISSUE] DOCX contained no readable paragraph/table content.")
    return "\n".join(lines).strip(), 1, issues, segments


def extract_and_prepare(file_path: str, filename: str) -> ExtractionResult:
    file_type = _detect_file_type(file_path, filename)
    issues: list[str] = []
    ocr_conf: float | None = None

    if file_type == "pdf":
        raw_text, total_pages, pdf_issues, ocr_conf, segments = _extract_pdf(file_path)
        issues.extend(pdf_issues)
    elif file_type == "image":
        raw_text, image_issues, ocr_conf, segments = _extract_image(file_path)
        total_pages = 1
        issues.extend(image_issues)
    elif file_type == "docx":
        raw_text, total_pages, docx_issues, segments = _extract_docx(file_path)
        issues.extend(docx_issues)
    else:
        raise ValueError("Unsupported file type. Supported: PDF, JPG, PNG, DOCX.")

    cleaned_text = _clean_text(raw_text)
    word_count = len(cleaned_text.split())
    language = _safe_detect_lang(cleaned_text)
    low_quality = len(cleaned_text) < 100

    if low_quality:
        issues.append("POOR_QUALITY_SCAN: Extraction yielded fewer than 100 characters.")
    if ocr_conf is not None and ocr_conf < 60:
        issues.append("LOW_OCR_CONFIDENCE: OCR confidence is below 60%. Upload a higher-resolution document.")
    if word_count < 50:
        issues.append("LOW_WORD_COUNT: Extracted text has fewer than 50 words.")

    red_flag_terms = [term for term in RED_FLAG_TERMS if re.search(term, cleaned_text, flags=re.IGNORECASE)]
    chunks = _split_chunks_preserve_paragraphs(cleaned_text)

    return ExtractionResult(
        file_type=file_type,
        total_pages=total_pages,
        raw_text=raw_text,
        cleaned_text=cleaned_text,
        extraction_issues=issues,
        low_quality_scan=low_quality,
        ocr_confidence=ocr_conf,
        language=language,
        word_count=word_count,
        red_flag_terms=red_flag_terms,
        chunks=chunks,
        segments=segments,
    )


def _extract_defined_terms(text: str) -> list[dict[str, str]]:
    terms: list[dict[str, str]] = []

    qpat = re.finditer(r"\"([A-Za-z][^\"\n]{1,60})\"\s+(?:means|shall mean)\s+([^.\n]{3,220})", text, flags=re.IGNORECASE)
    for m in qpat:
        terms.append({"term": m.group(1).strip(), "definition": m.group(2).strip()})

    cpat = re.finditer(r"\b([A-Z][A-Z\s]{2,40})\b\s+(?:means|shall mean)\s+([^.\n]{3,220})", text)
    for m in cpat:
        term = re.sub(r"\s+", " ", m.group(1)).strip()
        if 2 <= len(term.split()) <= 5:
            terms.append({"term": term, "definition": m.group(2).strip()})

    dedup = []
    seen = set()
    for t in terms:
        key = (t["term"].lower(), t["definition"].lower())
        if key in seen:
            continue
        seen.add(key)
        dedup.append(t)
    return dedup[:50]


def _extract_clauses(text: str) -> list[dict[str, str]]:
    lines = text.splitlines()
    clauses: list[dict[str, str]] = []
    heading_idx = []
    for idx, ln in enumerate(lines):
        s = ln.strip()
        if re.match(r"^\d+(\.\d+)*\s+[A-Za-z]", s) or re.match(r"^(Article|Section|Clause)\s+\w+", s, flags=re.IGNORECASE):
            heading_idx.append((idx, s))

    for i, (idx, head) in enumerate(heading_idx):
        end = heading_idx[i + 1][0] if i + 1 < len(heading_idx) else min(len(lines), idx + 8)
        body = " ".join(x.strip() for x in lines[idx + 1 : end] if x.strip())
        if not body:
            body = "[EXTRACTION_ISSUE] Clause body not clearly extracted."
        clauses.append({"clause": head, "what_it_says": body[:500], "risk_flag": ""})
    return clauses[:40]


def _extract_parties(text: str) -> list[dict[str, str]]:
    parties: list[dict[str, str]] = []
    patterns = [
        (r"\bbetween\s+(.+?)\s+and\s+(.+?)(?:\.|\n)", "Party"),
        (r"\b(Plaintiff|Defendant|Employer|Employee|Licensor|Licensee|Lessor|Lessee)\s*:\s*([^\n]+)", "Role"),
    ]
    for patt, mode in patterns:
        for m in re.finditer(patt, text, flags=re.IGNORECASE):
            if mode == "Party":
                parties.append({"name": m.group(1).strip(), "role": "Party A"})
                parties.append({"name": m.group(2).strip(), "role": "Party B"})
            else:
                parties.append({"name": m.group(2).strip(), "role": m.group(1).strip()})
    dedup = []
    seen = set()
    for p in parties:
        key = (p["name"].lower(), p["role"].lower())
        if key in seen:
            continue
        seen.add(key)
        dedup.append(p)
    return dedup[:30]


def _find_document_title(text: str) -> str:
    for line in text.splitlines()[:25]:
        s = line.strip()
        if len(s) < 5:
            continue
        if len(s) <= 140:
            return s
    return "Untitled Document"


def _guess_document_type(text: str) -> str:
    t = text.lower()
    if "non-disclosure" in t or "nda" in t:
        return "NDA"
    if "agreement" in t:
        return "Agreement"
    if "plaintiff" in t or "defendant" in t or "petition" in t:
        return "Court Filing"
    if "deed" in t:
        return "Deed"
    if "contract" in t:
        return "Contract"
    return "Other"


def _json_from_llm(raw: str) -> dict | None:
    txt = (raw or "").strip()
    try:
        obj = json.loads(txt)
        return obj if isinstance(obj, dict) else None
    except Exception:
        pass
    m = re.search(r"\{.*\}", txt, flags=re.DOTALL)
    if m:
        try:
            obj = json.loads(m.group(0))
            return obj if isinstance(obj, dict) else None
        except Exception:
            return None
    return None


def _fallback_structured(ex: ExtractionResult, filename: str) -> dict[str, Any]:
    text = ex.cleaned_text
    title = _find_document_title(text)
    dates = list(dict.fromkeys(DATE_REGEX.findall(text)))[:40]
    money = list(dict.fromkeys(MONEY_REGEX.findall(text)))[:40]
    parties = _extract_parties(text)
    defined_terms = _extract_defined_terms(text)
    clauses = _extract_clauses(text)

    critical = []
    for c in clauses:
        risk = ""
        lc = c["clause"].lower() + " " + c["what_it_says"].lower()
        if any(k in lc for k in ["terminate", "termination", "without notice", "indemn"]):
            risk = "Potential one-sided or high-risk obligation. Review carefully."
        critical.append(
            {
                "clause_name_number": c["clause"],
                "what_it_says": c["what_it_says"],
                "legal_implication_or_risk": risk,
            }
        )

    red_flags = []
    for term in ex.red_flag_terms:
        red_flags.append(
            f"Detected keyword '{term}' in extracted text. Manual legal review recommended."
        )
    if ex.low_quality_scan:
        red_flags.append("POOR_QUALITY_SCAN detected; important clauses may be missing.")
    if ex.ocr_confidence is not None and ex.ocr_confidence < 60:
        red_flags.append("OCR confidence below 60%. Upload a higher-resolution scan.")

    attachments = []
    for m in re.finditer(r"\b(?:Exhibit|Schedule|Annexure|Annex|Appendix)\s+[A-Za-z0-9-]+\b", text, flags=re.IGNORECASE):
        attachments.append(m.group(0))

    section_text = text[:1800]
    plain_summary = (
        f"This { _guess_document_type(text).lower() } appears to define legal obligations between the listed parties. "
        "The extracted content includes clauses that should be reviewed for payment terms, duration, termination, "
        "liability, confidentiality, dispute resolution, and governing law. "
        f"The document contains {len(dates)} date references and {len(money)} monetary references in extracted text. "
        "Any flagged extraction issues or low OCR confidence can affect legal interpretation and should be verified against the original file."
    )

    return {
        "document_overview": {
            "document_type": _guess_document_type(text),
            "title": title,
            "date": dates[0] if dates else "",
            "jurisdiction": "",
            "total_pages": str(ex.total_pages),
        },
        "parties_involved": parties,
        "purpose_scope": section_text[:600] if section_text else "Unable to reliably infer purpose from extracted text.",
        "key_terms_definitions": defined_terms,
        "critical_clauses": critical[:20],
        "red_flags_risk_analysis": red_flags,
        "important_dates_deadlines": dates,
        "financial_obligations": money,
        "obligations_summary": {
            "party_a_must": [],
            "party_b_must": [],
        },
        "attachments_exhibits": list(dict.fromkeys(attachments))[:30],
        "plain_english_summary": plain_summary,
        "disclaimer": "This analysis is AI-generated and for informational purposes only. It does not constitute legal advice. Consult a qualified attorney for legal decisions.",
    }


def _llm_structured_summary(
    ex: ExtractionResult,
    filename: str,
    llm_text_fn: Callable[[str], str] | None,
) -> dict[str, Any]:
    if llm_text_fn is None:
        return _fallback_structured(ex, filename)

    chunk_summaries = []
    for idx, chunk in enumerate(ex.chunks, start=1):
        prompt = (
            f"{SYSTEM_PROMPT}\n\n"
            f"Chunk {idx}/{len(ex.chunks)} from file {filename}:\n\n{chunk[:24000]}\n\n"
            "Return JSON with keys: chunk_summary, dates, amounts, parties, terms, clauses, extraction_issues."
        )
        try:
            raw = llm_text_fn(prompt)
            parsed = _json_from_llm(raw)
            if parsed:
                chunk_summaries.append(parsed)
        except Exception:
            continue

    if not chunk_summaries:
        return _fallback_structured(ex, filename)

    merge_prompt = (
        f"{SYSTEM_PROMPT}\n\n"
        f"Document metadata: file={filename}, pages={ex.total_pages}, language={ex.language}\n"
        f"Extraction issues: {ex.extraction_issues}\n"
        f"Red flag terms: {ex.red_flag_terms}\n"
        "Combine these chunk summaries into one final strict JSON object in the required schema.\n\n"
        f"{json.dumps(chunk_summaries, ensure_ascii=False)[:50000]}"
    )
    try:
        merged = _json_from_llm(llm_text_fn(merge_prompt))
        if isinstance(merged, dict):
            return merged
    except Exception:
        pass

    return _fallback_structured(ex, filename)


def _markdown_report(report: dict[str, Any], ex: ExtractionResult) -> str:
    overview = report.get("document_overview", {})
    parties = report.get("parties_involved", [])
    key_terms = report.get("key_terms_definitions", [])
    clauses = report.get("critical_clauses", [])
    red_flags = report.get("red_flags_risk_analysis", [])
    dates = report.get("important_dates_deadlines", [])
    money = report.get("financial_obligations", [])
    obligations = report.get("obligations_summary", {})
    attachments = report.get("attachments_exhibits", [])

    md = []
    if ex.low_quality_scan:
        md.append("**Note:** Document appears scanned/low quality. OCR quality may affect legal accuracy.\n")

    if ex.extraction_issues:
        md.append("**Extraction Issues:**")
        for issue in ex.extraction_issues:
            md.append(f"- {issue}")
        md.append("")

    md.append("## 📄 Document Overview")
    md.append(f"- Document Type: {overview.get('document_type', 'Other')}")
    md.append(f"- Title: {overview.get('title', '')}")
    md.append(f"- Date: {overview.get('date', '')}")
    md.append(f"- Jurisdiction: {overview.get('jurisdiction', '')}")
    md.append(f"- Total Pages: {overview.get('total_pages', '')}")
    md.append("")

    md.append("## 👥 Parties Involved")
    if parties:
        for p in parties:
            md.append(f"- {p.get('name', '')} — {p.get('role', '')}")
    else:
        md.append("- Not clearly identified in extracted text.")
    md.append("")

    md.append("## 🎯 Purpose & Scope")
    md.append(report.get("purpose_scope", ""))
    md.append("")

    md.append("## 📌 Key Terms & Definitions")
    if key_terms:
        for t in key_terms:
            md.append(f"- {t.get('term', '')}: {t.get('definition', '')}")
    else:
        md.append("- No explicit defined terms confidently extracted.")
    md.append("")

    md.append("## ⚖️ Critical Clauses")
    if clauses:
        for c in clauses:
            md.append(f"- Clause Name / Number: {c.get('clause_name_number', '')}")
            md.append(f"  What it says: {c.get('what_it_says', '')}")
            md.append(f"  Legal implication or risk flag: {c.get('legal_implication_or_risk', '')}")
    else:
        md.append("- No critical clauses confidently extracted.")
    md.append("")

    md.append("## 🚨 Red Flags & Risk Analysis")
    if red_flags:
        for r in red_flags:
            md.append(f"- {r}")
    else:
        md.append("- No explicit high-risk patterns detected from extracted text.")
    md.append("")

    md.append("## 📅 Important Dates & Deadlines")
    if dates:
        for d in dates:
            md.append(f"- {d}")
    else:
        md.append("- None clearly extracted.")
    md.append("")

    md.append("## 💰 Financial Obligations")
    if money:
        for m in money:
            md.append(f"- {m}")
    else:
        md.append("- None clearly extracted.")
    md.append("")

    md.append("## ✅ Obligations Summary")
    md.append("- Party A must:")
    for item in obligations.get("party_a_must", []):
        md.append(f"  - {item}")
    if not obligations.get("party_a_must"):
        md.append("  - Not clearly extracted.")
    md.append("- Party B must:")
    for item in obligations.get("party_b_must", []):
        md.append(f"  - {item}")
    if not obligations.get("party_b_must"):
        md.append("  - Not clearly extracted.")
    md.append("")

    md.append("## 📎 Attachments & Exhibits")
    if attachments:
        for a in attachments:
            md.append(f"- {a}")
    else:
        md.append("- None explicitly referenced in extracted text.")
    md.append("")

    md.append("## 💡 Plain-English Summary")
    md.append(report.get("plain_english_summary", ""))
    md.append("")

    md.append("## ⚠️ Disclaimer")
    md.append(
        report.get(
            "disclaimer",
            "This analysis is AI-generated and for informational purposes only. It does not constitute legal advice. Consult a qualified attorney for legal decisions.",
        )
    )

    return "\n".join(md).strip()


def _export_pdf(markdown_text: str, out_path: str) -> str:
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab is not installed; PDF export is unavailable.")
    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4
    text = c.beginText(40, height - 40)
    text.setFont("Helvetica", 10)
    for line in markdown_text.splitlines():
        wrapped = [line[i : i + 105] for i in range(0, len(line), 105)] or [""]
        for wline in wrapped:
            text.textLine(wline)
            if text.getY() < 40:
                c.drawText(text)
                c.showPage()
                text = c.beginText(40, height - 40)
                text.setFont("Helvetica", 10)
    c.drawText(text)
    c.save()
    return out_path


def _export_docx(markdown_text: str, out_path: str) -> str:
    doc = Document()
    for line in markdown_text.splitlines():
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(out_path)
    return out_path


def build_legal_document_report(
    file_path: str,
    filename: str,
    llm_text_fn: Callable[[str], str] | None = None,
    export_pdf: bool = False,
    export_docx: bool = False,
) -> dict[str, Any]:
    ex = extract_and_prepare(file_path, filename)
    structured = _llm_structured_summary(ex, filename, llm_text_fn)
    markdown = _markdown_report(structured, ex)

    # Ensure all required top-level report sections exist.
    required_keys = [
        "document_overview",
        "parties_involved",
        "purpose_scope",
        "key_terms_definitions",
        "critical_clauses",
        "red_flags_risk_analysis",
        "important_dates_deadlines",
        "financial_obligations",
        "obligations_summary",
        "attachments_exhibits",
        "plain_english_summary",
        "disclaimer",
    ]
    missing = [k for k in required_keys if k not in structured]
    for k in missing:
        structured[k] = [] if k.endswith("s") else ""

    exports: dict[str, str] = {}
    if export_pdf:
        pdf_path = os.path.join(tempfile.gettempdir(), f"legal_doc_summary_{os.getpid()}.pdf")
        exports["pdf_path"] = _export_pdf(markdown, pdf_path)
    if export_docx:
        docx_path = os.path.join(tempfile.gettempdir(), f"legal_doc_summary_{os.getpid()}.docx")
        exports["docx_path"] = _export_docx(markdown, docx_path)

    return {
        "status": "ok",
        "filename": filename,
        "file_type": ex.file_type,
        "quality_checks": {
            "word_count": ex.word_count,
            "language": ex.language,
            "low_quality_scan": ex.low_quality_scan,
            "ocr_confidence": ex.ocr_confidence,
            "red_flag_terms": ex.red_flag_terms,
            "extraction_issues": ex.extraction_issues,
            "chunks_used": len(ex.chunks),
            "requires_human_review": len(ex.red_flag_terms) >= 3,
            "is_large_document": ex.total_pages > 50,
        },
        "extracted_text": {
            "total_pages": ex.total_pages,
            "cleaned_text": ex.cleaned_text,
        },
        "report_json": structured,
        "report_markdown": markdown,
        "exports": exports,
    }


# ---------------------------------------------------------------------------
# Exhaustive document intelligence builder
# ---------------------------------------------------------------------------

CLAUSE_HEADING_RE = re.compile(
    r"^(?:(?:Article|Section|Clause|Schedule|Annexure|Annex|Appendix)\s+[A-Za-z0-9()./-]+|"
    r"\d+(?:\.\d+)*\s+[A-Z].+|"
    r"[IVXLC]+\.\s+[A-Z].+)$"
)

PARTY_ROLE_RE = re.compile(
    r"\b(Plaintiff|Defendant|Petitioner|Respondent|Employer|Employee|Licensor|Licensee|Witness|Attorney|Judge|Appellant|Respondent)\s*[:\-]\s*([^\n;,]+)",
    flags=re.IGNORECASE,
)

LEGAL_TERM_RULES = [
    ("Summons", re.compile(r"\bsummons?\b", re.IGNORECASE)),
    ("Subpoena", re.compile(r"\bsubpoena\b", re.IGNORECASE)),
    ("Injunction", re.compile(r"\binjunction\b", re.IGNORECASE)),
    ("Stay Order", re.compile(r"\bstay order\b", re.IGNORECASE)),
    ("Ex-Parte", re.compile(r"\bex[-\s]?parte\b", re.IGNORECASE)),
    ("Contempt", re.compile(r"\bcontempt\b", re.IGNORECASE)),
    ("Writ", re.compile(r"\bwrit\b", re.IGNORECASE)),
    ("Affidavit", re.compile(r"\baffidavit\b", re.IGNORECASE)),
    ("Deposition", re.compile(r"\bdeposition\b", re.IGNORECASE)),
    ("Discovery", re.compile(r"\bdiscovery\b", re.IGNORECASE)),
    ("Judgment", re.compile(r"\bjudgment\b", re.IGNORECASE)),
    ("Decree", re.compile(r"\bdecree\b", re.IGNORECASE)),
    ("Appellant", re.compile(r"\bappellant\b", re.IGNORECASE)),
    ("Respondent", re.compile(r"\brespondent\b", re.IGNORECASE)),
    ("Indemnification", re.compile(r"\bindemnif\w*\b", re.IGNORECASE)),
    ("Arbitration", re.compile(r"\barbitration\b", re.IGNORECASE)),
    ("Force Majeure", re.compile(r"\bforce majeure\b", re.IGNORECASE)),
    ("Liquidated Damages", re.compile(r"\bliquidated damages\b", re.IGNORECASE)),
    ("Consideration", re.compile(r"\bconsideration\b", re.IGNORECASE)),
    ("Novation", re.compile(r"\bnovation\b", re.IGNORECASE)),
    ("Escrow", re.compile(r"\bescrow\b", re.IGNORECASE)),
    ("Waiver", re.compile(r"\bwaiver\b", re.IGNORECASE)),
    ("Lien", re.compile(r"\blien\b", re.IGNORECASE)),
    ("Non-Disclosure", re.compile(r"\bnon[-\s]?disclosure\b|\bconfidentiality\b", re.IGNORECASE)),
    ("Non-Compete", re.compile(r"\bnon[-\s]?compete\b", re.IGNORECASE)),
    ("Termination for Cause", re.compile(r"\btermination for cause\b", re.IGNORECASE)),
    ("Intellectual Property", re.compile(r"\bintellectual property\b|\bIP\b", re.IGNORECASE)),
]

DATE_VALUE_RE = re.compile(
    r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|"
    r"\d{1,2}\s+(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{2,4}|"
    r"(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{1,2},\s*\d{2,4})\b",
    flags=re.IGNORECASE,
)

MONEY_VALUE_RE = re.compile(r"(?:₹\s?[\d,]+(?:\.\d{1,2})?|INR\s?[\d,]+(?:\.\d{1,2})?|USD\s?\$?[\d,]+(?:\.\d{1,2})?|\$[\d,]+(?:\.\d{1,2})?)")


def _seg_text(segment: dict[str, Any]) -> str:
    return str(segment.get("text") or "").strip()


def _seg_locator(segment: dict[str, Any]) -> str:
    if "page" in segment and segment["page"] not in (None, ""):
        return f"Page {segment['page']}"
    if "paragraph" in segment:
        return f"DOCX paragraph {segment['paragraph']}"
    if "table" in segment:
        return f"DOCX table {segment.get('table')} row {segment.get('row')}"
    return "Unknown locator"


def _segment_units(ex: ExtractionResult) -> list[dict[str, Any]]:
    units = [s for s in ex.segments if _seg_text(s)]
    if units:
        return units
    return [{"page": 1, "text": ex.cleaned_text or ex.raw_text or "", "source": "fallback"}]


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _sentence_windows(text: str) -> list[str]:
    if not text.strip():
        return []
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9(])", _norm(text))
    return [p.strip() for p in parts if p.strip()]


def _find_clause_heading(line: str) -> bool:
    return bool(CLAUSE_HEADING_RE.match(line.strip()))


def _extract_clauses_exhaustive(ex: ExtractionResult) -> list[dict[str, Any]]:
    clauses: list[dict[str, Any]] = []
    seen = set()

    for unit in _segment_units(ex):
        text = _seg_text(unit)
        lines = [ln.rstrip() for ln in text.splitlines()]
        current_heading = None
        current_body: list[str] = []

        def flush():
            nonlocal current_heading, current_body
            if not current_heading:
                return
            body = _norm(" ".join(current_body))
            exact_text = _norm("\n".join([current_heading] + current_body))
            key = (current_heading.lower(), exact_text[:250].lower())
            if key in seen:
                current_heading = None
                current_body = []
                return
            seen.add(key)
            clause_no_match = re.match(
                r"^(?:(?:Article|Section|Clause|Schedule|Annexure|Annex|Appendix)\s+[A-Za-z0-9()./-]+|\d+(?:\.\d+)*|[IVXLC]+)",
                current_heading,
                flags=re.IGNORECASE,
            )
            clause_no = clause_no_match.group(0).strip() if clause_no_match else current_heading.split(" ", 1)[0]
            clauses.append(
                {
                    "clause_no": clause_no,
                    "heading": current_heading,
                    "page": _seg_locator(unit),
                    "exact_text": exact_text,
                    "what_it_says": body or "NOT FOUND IN DOCUMENT",
                    "important": "YES" if any(k in (current_heading + " " + body).lower() for k in [
                        "payment", "term", "terminate", "indemn", "confidential", "arbitrat", "jurisdiction", "law", "ip", "non-compete", "notice"
                    ]) else "NO",
                }
            )
            current_heading = None
            current_body = []

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                if current_heading:
                    current_body.append("")
                continue
            if _find_clause_heading(line):
                flush()
                current_heading = line
                current_body = []
            else:
                if current_heading:
                    current_body.append(line)

        flush()

    if not clauses:
        return [
            {
                "clause_no": "NOT FOUND IN DOCUMENT",
                "heading": "NOT FOUND IN DOCUMENT",
                "page": "NOT FOUND IN DOCUMENT",
                "exact_text": "NOT FOUND IN DOCUMENT",
                "what_it_says": "NOT FOUND IN DOCUMENT",
                "important": "NO",
            }
        ]
    return clauses


def _extract_document_identity_exhaustive(ex: ExtractionResult, filename: str) -> dict[str, Any]:
    text = ex.cleaned_text or ex.raw_text or ""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    title = lines[0] if lines else filename
    for line in lines[:10]:
        if len(line) >= 5 and len(line) <= 160 and not line.lower().startswith(("page ", "table of contents")):
            title = line
            break

    case_number = ""
    case_patterns = [
        r"\b(?:Case|C\.?R\.?|Reference|Ref\.?|Suit|Appeal|Petition|Writ)\s*(?:No\.?|Number|No|#)?\s*[:\-]?\s*([A-Za-z0-9./()\-]+(?:\s*of\s*\d{2,4})?)",
        r"\b(?:No\.?|Number)\s*[:\-]?\s*([A-Za-z0-9./()\-]+)",
    ]
    for patt in case_patterns:
        m = re.search(patt, text, flags=re.IGNORECASE)
        if m:
            case_number = m.group(1).strip()
            break

    court_name = ""
    court_patterns = [
        r"\b([A-Z][A-Za-z&\s]+Court[^\n,]*)",
        r"\b([A-Z][A-Za-z&\s]+Tribunal[^\n,]*)",
        r"\b([A-Z][A-Za-z&\s]+Commission[^\n,]*)",
    ]
    for patt in court_patterns:
        m = re.search(patt, text, flags=re.IGNORECASE)
        if m:
            court_name = _norm(m.group(1))
            break

    jurisdiction = ""
    jur_patterns = [
        r"(?:governed by|subject to|jurisdiction of|exclusive jurisdiction of)\s+([^.\n]{2,140})",
        r"\bgoverning law\s*[:\-]\s*([^.\n]{2,140})",
    ]
    for patt in jur_patterns:
        m = re.search(patt, text, flags=re.IGNORECASE)
        if m:
            jurisdiction = _norm(m.group(1))
            break

    filing_date = ""
    for patt in [
        r"(?:dated|executed|filed|issued|made on|this\s+\d{1,2}(?:st|nd|rd|th)?\s+day\s+of)\s+([A-Za-z0-9,/\-\s]+)",
        DATE_VALUE_RE.pattern,
    ]:
        m = re.search(patt, text, flags=re.IGNORECASE)
        if m:
            filing_date = _norm(m.group(1 if m.lastindex else 0))
            break

    if not filing_date:
        date_hits = DATE_VALUE_RE.findall(text)
        filing_date = date_hits[0] if date_hits else ""

    language = ex.language or "unknown"
    doc_type = _guess_document_type(text)
    if any(k in text.lower() for k in ["summons", "affidavit", "petition", "writ", "order"]):
        if "summons" in text.lower():
            doc_type = "Summons"
        elif "affidavit" in text.lower():
            doc_type = "Affidavit"
        elif "petition" in text.lower():
            doc_type = "Petition"
        elif "writ" in text.lower():
            doc_type = "Writ"
        elif "order" in text.lower():
            doc_type = "Court Order"

    return {
        "document_title": title,
        "document_type": doc_type,
        "case_number_reference": case_number or "NOT FOUND IN DOCUMENT",
        "filing_execution_date": filing_date or "NOT FOUND IN DOCUMENT",
        "court_name_location": court_name or "NOT FOUND IN DOCUMENT",
        "jurisdiction_governing_law": jurisdiction or "NOT FOUND IN DOCUMENT",
        "total_pages": ex.total_pages,
        "language": language,
    }


def _extract_parties_exhaustive(ex: ExtractionResult) -> list[dict[str, Any]]:
    text = ex.cleaned_text or ex.raw_text or ""
    parties: list[dict[str, Any]] = []
    seen = set()

    def add_party(name: str, role: str = "Other", locator: str = "", address: str = "", reg_id: str = ""):
        clean = _norm(name)
        if not clean or len(clean) < 2:
            return
        key = (clean.lower(), role.lower(), address.lower(), reg_id.lower())
        if key in seen:
            return
        seen.add(key)
        parties.append(
            {
                "name": clean,
                "role": role,
                "address": address or "NOT FOUND IN DOCUMENT",
                "id_reg_no": reg_id or "NOT FOUND IN DOCUMENT",
                "locator": locator or "NOT FOUND IN DOCUMENT",
            }
        )

    # Party blocks introduced by "between ... and ..."
    for m in re.finditer(r"\bbetween\s+(.+?)\s+and\s+(.+?)(?:\n|\.|,)", text, flags=re.IGNORECASE | re.DOTALL):
        add_party(m.group(1), "Party A")
        add_party(m.group(2), "Party B")

    # Explicit role labels
    for m in PARTY_ROLE_RE.finditer(text):
        add_party(m.group(2), m.group(1).title())

    # Entity / person-style mentions in signature blocks or party definitions
    entity_patterns = [
        r"\b(?:M/s\.?|Mr\.?|Ms\.?|Mrs\.?|Dr\.?)\s+[A-Z][A-Za-z0-9&.,'()/-]+(?:\s+[A-Z0-9][A-Za-z0-9&.,'()/-]+){0,6}",
        r"\b[A-Z][A-Za-z0-9&.,'()/-]+(?:\s+[A-Z0-9][A-Za-z0-9&.,'()/-]+){1,6}\s+(?:Ltd\.?|Limited|Pvt\.?|Private Limited|LLP|LLC|Inc\.?|Company|Corporation)\b",
    ]
    for patt in entity_patterns:
        for m in re.finditer(patt, text):
            add_party(m.group(0), "Other")

    # Addresses / IDs keyed off nearby text
    address_match = re.search(r"(?:address|residing at|registered office at|office at)\s*[:\-]?\s*([^\n]{5,180})", text, flags=re.IGNORECASE)
    reg_match = re.search(r"(?:CIN|GSTIN|PAN|Aadhaar|Passport|Registration No\.?|Reg(?:istration)? No\.?)\s*[:\-]?\s*([A-Za-z0-9./\-]+)", text, flags=re.IGNORECASE)
    address = _norm(address_match.group(1)) if address_match else ""
    reg_id = _norm(reg_match.group(1)) if reg_match else ""

    for p in parties:
        if p["address"] == "NOT FOUND IN DOCUMENT" and address:
            p["address"] = address
        if p["id_reg_no"] == "NOT FOUND IN DOCUMENT" and reg_id:
            p["id_reg_no"] = reg_id

    if not parties:
        return [
            {
                "name": "NOT FOUND IN DOCUMENT",
                "role": "NOT FOUND IN DOCUMENT",
                "address": "NOT FOUND IN DOCUMENT",
                "id_reg_no": "NOT FOUND IN DOCUMENT",
                "locator": "NOT FOUND IN DOCUMENT",
            }
        ]

    return parties[:80]


def _extract_critical_alerts_exhaustive(ex: ExtractionResult, clauses: list[dict[str, Any]]) -> dict[str, Any]:
    text = ex.cleaned_text or ex.raw_text or ""
    units = _segment_units(ex)
    alerts = {
        "summons": [],
        "court_hearings": [],
        "notices": [],
        "orders_directions": [],
        "warrants": [],
        "injunctions": [],
        "appeals": [],
    }

    def nearest_clause(context: str) -> dict[str, Any]:
        context_l = context.lower()
        for clause in clauses:
            if clause.get("exact_text") and any(tok in clause["exact_text"].lower() for tok in context_l.split()[:8]):
                return clause
        return {}

    for unit in units:
        utext = _seg_text(unit)
        locator = _seg_locator(unit)
        lower = utext.lower()
        if "summons" in lower:
            snippet = _norm(utext)
            alerts["summons"].append(
                {
                    "locator": locator,
                    "exact_text": snippet[:600],
                    "plain_english": "A summons is present and requires the named person or entity to respond or appear as directed.",
                    "clause": nearest_clause(snippet).get("clause_no", "NOT FOUND IN DOCUMENT"),
                    "deadline_to_respond": _extract_first_date_near_text(snippet) or "NOT FOUND IN DOCUMENT",
                }
            )

        if any(k in lower for k in ["hearing", "listed on", "returnable", "next date", "date of hearing"]):
            for d in DATE_VALUE_RE.finditer(utext):
                window = _norm(utext[max(0, d.start()-90): min(len(utext), d.end()+140)])
                alerts["court_hearings"].append(
                    {
                        "hearing_date": d.group(0),
                        "time": _extract_time_near_text(window) or "NOT FOUND IN DOCUMENT",
                        "court": _extract_court_near_text(window) or "NOT FOUND IN DOCUMENT",
                        "purpose": _infer_purpose_from_context(window, "hearing"),
                        "locator": locator,
                        "exact_text": window,
                    }
                )

        if "notice" in lower or "show cause" in lower or "demand notice" in lower:
            alerts["notices"].append(
                {
                    "locator": locator,
                    "exact_text": _norm(utext)[:600],
                    "to_whom": _extract_party_like_target(utext) or "NOT FOUND IN DOCUMENT",
                    "reason": _infer_purpose_from_context(utext, "notice"),
                    "response_deadline": _extract_first_date_near_text(utext) or "NOT FOUND IN DOCUMENT",
                }
            )

        if any(k in lower for k in ["ordered", "directed", "shall", "is hereby", "hereby", "direction", "disposed"]):
            alerts["orders_directions"].append(
                {
                    "locator": locator,
                    "exact_text": _norm(utext)[:700],
                    "what_must_be_done": _infer_purpose_from_context(utext, "order"),
                    "by_when": _extract_first_date_near_text(utext) or "NOT FOUND IN DOCUMENT",
                }
            )

        if "warrant" in lower:
            alerts["warrants"].append(
                {
                    "locator": locator,
                    "exact_text": _norm(utext)[:600],
                    "type": _infer_warrant_type(utext),
                    "against_whom": _extract_party_like_target(utext) or "NOT FOUND IN DOCUMENT",
                    "reason": _infer_purpose_from_context(utext, "warrant"),
                }
            )

        if "injunction" in lower or "restrained" in lower or "prohibited" in lower:
            alerts["injunctions"].append(
                {
                    "locator": locator,
                    "exact_text": _norm(utext)[:700],
                    "what_is_restricted": _infer_restriction(utext),
                    "who_is_affected": _extract_party_like_target(utext) or "NOT FOUND IN DOCUMENT",
                    "duration": _extract_duration(utext) or "NOT FOUND IN DOCUMENT",
                }
            )

        if "appeal" in lower or "appellant" in lower or "respondent" in lower:
            alerts["appeals"].append(
                {
                    "locator": locator,
                    "exact_text": _norm(utext)[:700],
                    "appeal_against": _infer_purpose_from_context(utext, "appeal"),
                    "filed_by": _extract_party_like_target(utext) or "NOT FOUND IN DOCUMENT",
                    "deadline": _extract_first_date_near_text(utext) or "NOT FOUND IN DOCUMENT",
                }
            )

    for key in alerts:
        if not alerts[key]:
            alerts[key] = [{"note": "NOT FOUND IN DOCUMENT"}]
    return alerts


def _extract_all_dates_exhaustive(ex: ExtractionResult) -> list[dict[str, Any]]:
    text = ex.cleaned_text or ex.raw_text or ""
    seen = set()
    items: list[dict[str, Any]] = []

    for unit in _segment_units(ex):
        utext = _seg_text(unit)
        locator = _seg_locator(unit)
        for m in DATE_VALUE_RE.finditer(utext):
            date_text = m.group(0)
            context = _norm(utext[max(0, m.start()-100): min(len(utext), m.end()+160)])
            key = (date_text.lower(), context[:120].lower())
            if key in seen:
                continue
            seen.add(key)
            items.append(
                {
                    "date": date_text,
                    "what_it_is": _infer_date_type(context),
                    "who_it_affects": _infer_who_affects(context),
                    "locator": locator,
                    "exact_text": context,
                    "sort_key": _parse_date_sort_key(date_text),
                }
            )

    # Relative deadlines
    for m in re.finditer(r"\bwithin\s+(\d{1,3})\s+(day|days|week|weeks|month|months)\b", text, flags=re.IGNORECASE):
        context = _norm(text[max(0, m.start()-100): min(len(text), m.end()+160)])
        key = (context[:120].lower(), "relative")
        if key in seen:
            continue
        seen.add(key)
        items.append(
            {
                "date": m.group(0),
                "what_it_is": "Relative deadline",
                "who_it_affects": _infer_who_affects(context),
                "locator": "Document text",
                "exact_text": context,
                "sort_key": datetime.max,
            }
        )

    if not items:
        return [{"date": "NOT FOUND IN DOCUMENT", "what_it_is": "NOT FOUND IN DOCUMENT", "who_it_affects": "NOT FOUND IN DOCUMENT"}]

    items.sort(key=lambda x: x.get("sort_key", datetime.max))
    for item in items:
        item.pop("sort_key", None)
    return items


def _extract_financial_details_exhaustive(ex: ExtractionResult) -> list[dict[str, Any]]:
    text = ex.cleaned_text or ex.raw_text or ""
    items: list[dict[str, Any]] = []
    seen = set()
    for unit in _segment_units(ex):
        utext = _seg_text(unit)
        locator = _seg_locator(unit)
        for m in MONEY_VALUE_RE.finditer(utext):
            amount = m.group(0)
            context = _norm(utext[max(0, m.start()-110): min(len(utext), m.end()+180)])
            key = (amount.lower(), context[:120].lower())
            if key in seen:
                continue
            seen.add(key)
            items.append(
                {
                    "amount": amount,
                    "purpose": _infer_money_purpose(context),
                    "who_pays": _infer_who_affects(context),
                    "due_date": _extract_first_date_near_text(context) or "NOT FOUND IN DOCUMENT",
                    "locator": locator,
                    "exact_text": context,
                }
            )
    if not items:
        return [{"amount": "NOT FOUND IN DOCUMENT", "purpose": "NOT FOUND IN DOCUMENT", "who_pays": "NOT FOUND IN DOCUMENT", "due_date": "NOT FOUND IN DOCUMENT"}]
    return items


def _extract_legal_terms_exhaustive(ex: ExtractionResult, clauses: list[dict[str, Any]]) -> list[dict[str, Any]]:
    text = ex.cleaned_text or ex.raw_text or ""
    terms: list[dict[str, Any]] = []
    seen = set()
    for term_name, patt in LEGAL_TERM_RULES:
        for m in patt.finditer(text):
            context = _norm(text[max(0, m.start()-100): min(len(text), m.end()+180)])
            key = (term_name.lower(), context[:120].lower())
            if key in seen:
                continue
            seen.add(key)
            clause_ref = _match_clause_for_text(clauses, context)
            terms.append(
                {
                    "term": term_name,
                    "found_in": clause_ref.get("clause_no") or _locator_from_text(ex, context) or "NOT FOUND IN DOCUMENT",
                    "document_says": context,
                    "plain_english": _plain_term_explanation(term_name),
                }
            )
    if not terms:
        return [{"term": "NOT FOUND IN DOCUMENT", "found_in": "NOT FOUND IN DOCUMENT", "document_says": "NOT FOUND IN DOCUMENT", "plain_english": "NOT FOUND IN DOCUMENT"}]
    return terms


def _extract_red_flags_exhaustive(ex: ExtractionResult, clauses: list[dict[str, Any]]) -> list[dict[str, Any]]:
    text = ex.cleaned_text or ex.raw_text or ""
    flags: list[dict[str, Any]] = []
    try:
        analyzed = analyze_contract(text, None) if callable(analyze_contract) else {"flags": []}
    except Exception:
        analyzed = {"flags": []}
    for item in analyzed.get("flags", []):
        clause_ref = _match_clause_for_text(clauses, item.get("snippet", "") + " " + item.get("matched_text", ""))
        issue = {
            "risk_level": str(item.get("severity", "medium")).upper(),
            "clause": clause_ref.get("clause_no") or "NOT FOUND IN DOCUMENT",
            "issue": item.get("clause", "Potential risky clause pattern detected."),
            "impact": "May create one-sided or legally risky obligations.",
            "suggestion": f"Grounded by: {', '.join(c.get('title', '') for c in item.get('constitutional_grounding', [])) or 'Review with counsel'}",
            "exact_text": item.get("snippet") or item.get("matched_text") or "NOT FOUND IN DOCUMENT",
        }
        flags.append(issue)
    if ex.low_quality_scan:
        flags.append(
            {
                "risk_level": "HIGH",
                "clause": "EXTRACTION ISSUE",
                "issue": "Extraction yielded fewer than 100 characters.",
                "impact": "Important clauses may be missing from the scan.",
                "suggestion": "Upload a clearer scan or a text-based PDF.",
                "exact_text": "POOR_QUALITY_SCAN",
            }
        )
    if ex.ocr_confidence is not None and ex.ocr_confidence < 60:
        flags.append(
            {
                "risk_level": "HIGH",
                "clause": "EXTRACTION ISSUE",
                "issue": "OCR confidence below 60%.",
                "impact": "The OCR result may be unreliable.",
                "suggestion": "Use a higher-resolution upload.",
                "exact_text": f"OCR confidence: {ex.ocr_confidence}%",
            }
        )
    if not flags:
        return [{"risk_level": "NOTE", "clause": "NOT FOUND IN DOCUMENT", "issue": "No obvious risky clauses detected by the pattern scan.", "impact": "No immediate red flag identified from the extracted text.", "suggestion": "Still review the full document manually.", "exact_text": "NOT FOUND IN DOCUMENT"}]
    return flags


def _extract_obligations_exhaustive(ex: ExtractionResult, clauses: list[dict[str, Any]]) -> dict[str, Any]:
    text = ex.cleaned_text or ex.raw_text or ""
    party_map: dict[str, list[str]] = defaultdict(list)
    court_map: list[str] = []
    sentences = _sentence_windows(text)
    for sent in sentences:
        lower = sent.lower()
        if any(k in lower for k in ["shall", "must", "agrees to", "undertakes to", "is required to", "to be", "needs to"]):
            owner = _infer_party_owner(sent)
            if owner == "Court / Authority":
                court_map.append(sent)
            else:
                party_map[owner].append(sent)
    if not party_map:
        party_map["NOT FOUND IN DOCUMENT"] = ["NOT FOUND IN DOCUMENT"]
    return {
        "per_party": {k: v[:80] for k, v in party_map.items()},
        "court_authority": court_map[:40] or ["NOT FOUND IN DOCUMENT"],
    }


def _extract_attachments_exhaustive(ex: ExtractionResult) -> list[dict[str, Any]]:
    text = ex.cleaned_text or ex.raw_text or ""
    refs: list[dict[str, Any]] = []
    seen = set()
    for m in re.finditer(r"\b(?:Schedule|Schedule A|Schedule B|Exhibit|Exhibit 1|Annexure|Annex|Appendix)\s+[A-Za-z0-9-]*\b", text, flags=re.IGNORECASE):
        ref = _norm(m.group(0))
        key = ref.lower()
        if key in seen:
            continue
        seen.add(key)
        refs.append(
            {
                "reference": ref,
                "what_it_is": _attachment_guess(ref, text),
                "attached": "Yes" if ref.lower() in text.lower() else "No",
                "locator": _locator_from_text(ex, ref),
            }
        )
    if not refs:
        return [{"reference": "NOT FOUND IN DOCUMENT", "what_it_is": "NOT FOUND IN DOCUMENT", "attached": "NOT FOUND IN DOCUMENT"}]
    return refs


def _compose_plain_summary_exhaustive(report: dict[str, Any]) -> dict[str, Any]:
    identity = report.get("document_identity", {})
    alerts = report.get("critical_legal_alerts", {})
    dates = report.get("all_dates_deadlines", [])
    red_flags = report.get("red_flags_risks", [])
    parties = report.get("all_parties_involved", [])
    obligations = report.get("obligations", {})
    lines = []
    lines.append(f"This document appears to be a {identity.get('document_type', 'Other')} titled '{identity.get('document_title', 'NOT FOUND IN DOCUMENT')}'.")
    lines.append("It establishes legal obligations and contains the extracted parties, dates, clauses, financial items, and any priority legal alerts found in the text.")
    lines.append("The report below preserves exact quotations and marks missing items as NOT FOUND IN DOCUMENT.")
    return {
        "what_is_this_document": " ".join(lines[:2]),
        "what_is_happening": "The document has been scanned for parties, dates, clauses, financial obligations, legal terms, and priority alerts. Where the text is clear, exact text is preserved; where it is unclear, the output says so explicitly.",
        "what_each_party_needs_to_do": obligations.get("per_party", {}),
        "important_dates_to_remember": [d.get("date") for d in dates[:5] if isinstance(d, dict)],
        "what_happens_if_not_followed": "Non-compliance could create legal, financial, or procedural risk depending on the clause or order identified.",
        "should_i_be_worried": red_flags[:3] if red_flags else ["No explicit high-risk clause was detected, but manual review is still recommended."],
        "bottom_line": "Review the exact clauses and deadlines before acting; if anything is unclear, verify against the original document.",
        "party_count": len(parties),
        "priority_alerts_found": {
            "summons": len(alerts.get("summons", [])) if isinstance(alerts, dict) else 0,
            "hearings": len(alerts.get("court_hearings", [])) if isinstance(alerts, dict) else 0,
        },
    }


def _extract_first_date_near_text(text: str) -> str:
    m = DATE_VALUE_RE.search(text or "")
    return m.group(0) if m else ""


def _extract_time_near_text(text: str) -> str:
    m = re.search(r"\b\d{1,2}(?::\d{2})?\s*(?:AM|PM|am|pm)\b", text or "")
    return m.group(0) if m else ""


def _extract_court_near_text(text: str) -> str:
    m = re.search(r"\b[A-Z][A-Za-z&\s]+Court[^\n,;]*", text or "")
    return _norm(m.group(0)) if m else ""


def _infer_purpose_from_context(text: str, kind: str) -> str:
    low = (text or "").lower()
    if kind == "hearing":
        if "adjourn" in low:
            return "Adjourned hearing"
        if "interim" in low:
            return "Interim hearing"
        return "Court hearing"
    if kind == "notice":
        if "show cause" in low:
            return "Show-cause notice"
        if "demand" in low:
            return "Demand notice"
        return "Legal notice"
    if kind == "order":
        return "Court order or direction"
    if kind == "warrant":
        return "Warrant"
    if kind == "appeal":
        return "Appeal-related text"
    return "Legal reference"


def _extract_party_like_target(text: str) -> str:
    m = re.search(r"\b(?:against|to|upon|upon)\s+([A-Z][A-Za-z0-9&.,'()/-]+(?:\s+[A-Z0-9][A-Za-z0-9&.,'()/-]+){0,5})", text or "")
    return _norm(m.group(1)) if m else ""


def _infer_warrant_type(text: str) -> str:
    low = (text or "").lower()
    if "non-bailable" in low:
        return "Non-bailable warrant"
    if "bailable" in low:
        return "Bailable warrant"
    return "Warrant"


def _infer_restriction(text: str) -> str:
    low = (text or "").lower()
    if "restrained from" in low:
        return _norm(re.search(r"restrained from\s+([^.;\n]+)", text, flags=re.IGNORECASE).group(1)) if re.search(r"restrained from\s+([^.;\n]+)", text, flags=re.IGNORECASE) else _norm(text[:220])
    if "prohibited from" in low:
        return _norm(re.search(r"prohibited from\s+([^.;\n]+)", text, flags=re.IGNORECASE).group(1)) if re.search(r"prohibited from\s+([^.;\n]+)", text, flags=re.IGNORECASE) else _norm(text[:220])
    return _norm(text[:220])


def _extract_duration(text: str) -> str:
    m = re.search(r"\bfor\s+(?:a\s+period\s+of\s+)?([^.;\n]+)", text or "", flags=re.IGNORECASE)
    return _norm(m.group(1)) if m else ""


def _infer_date_type(context: str) -> str:
    low = (context or "").lower()
    if any(k in low for k in ["hearing", "listed on", "returnable", "adjourned"]):
        return "Hearing Date"
    if any(k in low for k in ["payment", "fee", "deposit", "installment", "invoice"]):
        return "Payment Due"
    if any(k in low for k in ["notice", "show cause", "reply", "response"]):
        return "Notice Response Deadline"
    if any(k in low for k in ["file", "submit", "within", "days", "weeks", "months"]):
        return "Deadline"
    if any(k in low for k in ["execution", "executed", "dated", "start"]):
        return "Execution / Start Date"
    return "Date Mentioned"


def _infer_who_affects(context: str) -> str:
    low = (context or "").lower()
    roles = []
    for role in ["plaintiff", "defendant", "petitioner", "respondent", "appellant", "licensor", "licensee", "employer", "employee", "party a", "party b"]:
        if role in low:
            roles.append(role.title())
    return ", ".join(dict.fromkeys(roles)) if roles else "Not clearly stated"


def _parse_date_sort_key(date_text: str):
    if not date_text:
        return datetime.max
    formats = [
        "%d/%m/%Y", "%d-%m-%Y", "%d/%m/%y", "%d-%m-%y",
        "%d %b %Y", "%d %B %Y", "%b %d, %Y", "%B %d, %Y",
        "%d %b %y", "%d %B %y", "%b %d, %y", "%B %d, %y",
    ]
    cleaned = date_text.replace("  ", " ").strip()
    for fmt in formats:
        try:
            return datetime.strptime(cleaned, fmt)
        except Exception:
            continue
    return datetime.max


def _infer_money_purpose(context: str) -> str:
    low = (context or "").lower()
    if "security deposit" in low:
        return "Security deposit"
    if "monthly" in low or "per month" in low:
        return "Monthly fee"
    if "penalty" in low or "liquidated" in low:
        return "Penalty / liquidated damages"
    if "interest" in low:
        return "Interest"
    if "salary" in low or "wages" in low:
        return "Salary / wages"
    if "fee" in low or "charges" in low:
        return "Fee / charges"
    return "Monetary obligation"


def _infer_party_owner(sentence: str) -> str:
    low = (sentence or "").lower()
    for role in ["party a", "party b", "plaintiff", "defendant", "petitioner", "respondent", "appellant", "licensor", "licensee", "employer", "employee", "court"]:
        if role in low:
            return role.title()
    return "Not clearly stated"


def _match_clause_for_text(clauses: list[dict[str, Any]], context: str) -> dict[str, Any]:
    ctx = (context or "").lower()
    for clause in clauses:
        text = (clause.get("exact_text") or "").lower()
        if not text:
            continue
        clause_no = str(clause.get("clause_no") or "")
        heading = str(clause.get("heading") or "")
        if clause_no and clause_no.lower() in ctx:
            return clause
        if heading and heading.lower()[:120] in ctx:
            return clause
        if any(tok for tok in re.findall(r"[a-zA-Z]{5,}", ctx[:160]) if tok in text):
            return clause
    return {}


def _locator_from_text(ex: ExtractionResult, snippet: str) -> str:
    snippet_l = (snippet or "").lower()
    for unit in _segment_units(ex):
        utext = _seg_text(unit)
        if snippet_l and snippet_l[:25] in utext.lower():
            return _seg_locator(unit)
    return "NOT FOUND IN DOCUMENT"


def _plain_term_explanation(term: str) -> str:
    mapping = {
        "Summons": "Official court notice requiring a person or party to appear or respond.",
        "Subpoena": "Order to produce documents or testify.",
        "Injunction": "Court order stopping or limiting an action.",
        "Stay Order": "Temporary pause on proceedings or enforcement.",
        "Ex-Parte": "Decision made without one party present.",
        "Contempt": "Disobeying a court order or obstructing the court.",
        "Writ": "Formal written court order.",
        "Affidavit": "Sworn written statement of facts.",
        "Deposition": "Out-of-court sworn testimony.",
        "Discovery": "Pre-trial exchange of information and evidence.",
        "Judgment": "Final decision of the court.",
        "Decree": "Court order with specific directions.",
        "Appellant": "Person who files an appeal.",
        "Respondent": "Person who responds to an appeal.",
        "Indemnification": "One party agrees to cover losses or damages of the other.",
        "Arbitration": "Private dispute resolution outside court.",
        "Force Majeure": "Protection for uncontrollable events such as floods or war.",
        "Liquidated Damages": "Pre-agreed amount payable on breach.",
        "Consideration": "What each party gives in exchange.",
        "Novation": "Replacing an old contract with a new one.",
        "Escrow": "Funds held by a third party until conditions are met.",
        "Waiver": "Giving up a legal right.",
        "Lien": "Legal right over property until debt is paid.",
        "Non-Disclosure": "Keeping information confidential.",
        "Non-Compete": "Restriction from working with competitors.",
        "Termination for Cause": "Ending a contract because of a specific violation.",
        "Intellectual Property": "Ownership of creations, inventions, or brands.",
    }
    return mapping.get(term, "Legal term detected in the document.")


def _attachment_guess(reference: str, text: str) -> str:
    low = (reference + " " + text).lower()
    if "schedule" in low:
        return "Schedule or attached schedule reference"
    if "exhibit" in low:
        return "Exhibit reference"
    if "annexure" in low or "annex" in low:
        return "Annexure reference"
    if "appendix" in low:
        return "Appendix reference"
    return "Referenced attachment or external document"


def _build_priority_alert_summary(alerts: dict[str, Any]) -> dict[str, Any]:
    summons = alerts.get("summons", []) if isinstance(alerts, dict) else []
    hearings = alerts.get("court_hearings", []) if isinstance(alerts, dict) else []
    return {
        "summons": summons,
        "court_hearings": hearings,
        "notices": alerts.get("notices", []) if isinstance(alerts, dict) else [],
        "orders_directions": alerts.get("orders_directions", []) if isinstance(alerts, dict) else [],
        "warrants": alerts.get("warrants", []) if isinstance(alerts, dict) else [],
        "injunctions": alerts.get("injunctions", []) if isinstance(alerts, dict) else [],
        "appeals": alerts.get("appeals", []) if isinstance(alerts, dict) else [],
    }


def _markdown_table(headers: list[str], rows: list[list[Any]]) -> str:
    if not rows:
        return "NOT FOUND IN DOCUMENT"
    header_line = " | ".join(headers)
    sep = " | ".join(["---"] * len(headers))
    body = "\n".join(" | ".join(str(cell) for cell in row) for row in rows)
    return f"{header_line}\n{sep}\n{body}"


def _markdown_from_exhaustive(report: dict[str, Any], ex: ExtractionResult) -> str:
    lines: list[str] = []
    identity = report["document_identity"]
    alerts = report["critical_legal_alerts"]
    parties = report["all_parties_involved"]
    dates = report["all_dates_deadlines"]
    money = report["all_financial_details"]
    terms = report["important_legal_terms_found"]
    clauses = report["all_clauses_full_list"]
    obligations = report["obligations"]
    risks = report["red_flags_risks"]
    attachments = report["attachments_exhibits_references"]
    summary = report["plain_english_summary"]

    lines.append("## PRIORITY ALERTS")
    lines.append("### SUMMONS")
    if alerts.get("summons") and alerts["summons"][0].get("note") != "NOT FOUND IN DOCUMENT":
        for item in alerts["summons"]:
            lines.append(f"- Locator: {item.get('locator', '')}")
            lines.append(f"- Exact text: {item.get('exact_text', '')}")
            lines.append(f"- What it means: {item.get('plain_english', '')}")
            lines.append(f"- Deadline to respond: {item.get('deadline_to_respond', '')}")
    else:
        lines.append("NOT FOUND IN DOCUMENT")
    lines.append("### HEARING DATES")
    if alerts.get("court_hearings") and alerts["court_hearings"][0].get("note") != "NOT FOUND IN DOCUMENT":
        for item in alerts["court_hearings"]:
            lines.append(f"- {item.get('hearing_date', '')} | {item.get('time', '')} | {item.get('court', '')} | {item.get('purpose', '')}")
            lines.append(f"  Exact text: {item.get('exact_text', '')}")
    else:
        lines.append("NOT FOUND IN DOCUMENT")
    lines.append("")

    lines.append("## SECTION 1 — 📄 DOCUMENT IDENTITY")
    rows = [
        ["Document Title", identity.get("document_title", "NOT FOUND IN DOCUMENT")],
        ["Document Type", identity.get("document_type", "NOT FOUND IN DOCUMENT")],
        ["Case Number / Reference Number", identity.get("case_number_reference", "NOT FOUND IN DOCUMENT")],
        ["Filing Date / Execution Date", identity.get("filing_execution_date", "NOT FOUND IN DOCUMENT")],
        ["Court Name & Location", identity.get("court_name_location", "NOT FOUND IN DOCUMENT")],
        ["Jurisdiction & Governing Law", identity.get("jurisdiction_governing_law", "NOT FOUND IN DOCUMENT")],
        ["Total Pages", identity.get("total_pages", "NOT FOUND IN DOCUMENT")],
        ["Language", identity.get("language", "NOT FOUND IN DOCUMENT")],
    ]
    lines.append(_markdown_table(["Field", "Value"], rows))
    lines.append("")

    lines.append("## SECTION 2 — 👥 ALL PARTIES INVOLVED")
    if parties and parties[0].get("name") != "NOT FOUND IN DOCUMENT":
        for p in parties:
            lines.append(f"- Name: {p.get('name', '')}")
            lines.append(f"  Role: {p.get('role', '')}")
            lines.append(f"  Address: {p.get('address', '')}")
            lines.append(f"  ID / Reg No: {p.get('id_reg_no', '')}")
            lines.append(f"  Locator: {p.get('locator', '')}")
    else:
        lines.append("NOT FOUND IN DOCUMENT")
    lines.append("")

    lines.append("## SECTION 3 — ⚠️ CRITICAL LEGAL ALERTS")
    for key in ["summons", "court_hearings", "notices", "orders_directions", "warrants", "injunctions", "appeals"]:
        lines.append(f"### {key.replace('_', ' ').upper()}")
        items = alerts.get(key, [])
        if items and items[0].get("note") != "NOT FOUND IN DOCUMENT":
            for item in items:
                for k, v in item.items():
                    lines.append(f"- {k}: {v}")
        else:
            lines.append("NOT FOUND IN DOCUMENT")
    lines.append("")

    lines.append("## SECTION 4 — 📅 ALL DATES & DEADLINES (COMPLETE LIST)")
    if dates and dates[0].get("date") != "NOT FOUND IN DOCUMENT":
        rows = [[d.get("date", ""), d.get("what_it_is", ""), d.get("who_it_affects", "")] for d in dates]
        lines.append(_markdown_table(["Date", "What It Is", "Who It Affects"], rows))
        for d in dates:
            lines.append(f"- Locator: {d.get('locator', '')}")
            lines.append(f"  Exact text: {d.get('exact_text', '')}")
    else:
        lines.append("NOT FOUND IN DOCUMENT")
    lines.append("")

    lines.append("## SECTION 5 — 💰 ALL FINANCIAL DETAILS (COMPLETE LIST)")
    if money and money[0].get("amount") != "NOT FOUND IN DOCUMENT":
        rows = [[m.get("amount", ""), m.get("purpose", ""), m.get("who_pays", ""), m.get("due_date", "")] for m in money]
        lines.append(_markdown_table(["Amount", "Purpose", "Who Pays", "Due Date"], rows))
        for m in money:
            lines.append(f"- Locator: {m.get('locator', '')}")
            lines.append(f"  Exact text: {m.get('exact_text', '')}")
    else:
        lines.append("NOT FOUND IN DOCUMENT")
    lines.append("")

    lines.append("## SECTION 6 — 📌 IMPORTANT LEGAL TERMS FOUND (GLOSSARY)")
    if terms and terms[0].get("term") != "NOT FOUND IN DOCUMENT":
        for t in terms:
            lines.append(f"- Term: {t.get('term', '')}")
            lines.append(f"  Found in: {t.get('found_in', '')}")
            lines.append(f"  Document says: {t.get('document_says', '')}")
            lines.append(f"  Plain English: {t.get('plain_english', '')}")
    else:
        lines.append("NOT FOUND IN DOCUMENT")
    lines.append("")

    lines.append("## SECTION 7 — 📋 ALL CLAUSES — FULL LIST")
    if clauses and clauses[0].get("clause_no") != "NOT FOUND IN DOCUMENT":
        for c in clauses:
            lines.append(f"- Clause No.: {c.get('clause_no', '')}")
            lines.append(f"  Heading: {c.get('heading', '')}")
            lines.append(f"  What it says: {c.get('what_it_says', '')}")
            lines.append(f"  Exact text: {c.get('exact_text', '')}")
            lines.append(f"  Important?: {c.get('important', '')}")
            lines.append(f"  Page/Locator: {c.get('page', '')}")
    else:
        lines.append("NOT FOUND IN DOCUMENT")
    lines.append("")

    lines.append("## SECTION 8 — ✅ OBLIGATIONS — WHO MUST DO WHAT")
    per_party = obligations.get("per_party", {}) if isinstance(obligations, dict) else {}
    if per_party:
        for party, items in per_party.items():
            lines.append(f"- {party}: must —")
            for item in items:
                lines.append(f"  □ {item}")
    else:
        lines.append("NOT FOUND IN DOCUMENT")
    court_items = obligations.get("court_authority", []) if isinstance(obligations, dict) else []
    lines.append("- COURT / AUTHORITY (if applicable): has directed —")
    if court_items:
        for item in court_items:
            lines.append(f"  □ {item}")
    else:
        lines.append("  □ NOT FOUND IN DOCUMENT")
    lines.append("")

    lines.append("## SECTION 9 — 🚨 RED FLAGS & RISKS (COMPLETE LIST)")
    if risks and risks[0].get("risk_level") != "NOTE" or (risks and risks[0].get("clause") != "NOT FOUND IN DOCUMENT"):
        for r in risks:
            lines.append(f"- Risk Level: {r.get('risk_level', '')}")
            lines.append(f"  Clause: {r.get('clause', '')}")
            lines.append(f"  Issue: {r.get('issue', '')}")
            lines.append(f"  Impact: {r.get('impact', '')}")
            lines.append(f"  Suggestion: {r.get('suggestion', '')}")
            lines.append(f"  Exact text: {r.get('exact_text', '')}")
    else:
        lines.append("NOT FOUND IN DOCUMENT")
    lines.append("")

    lines.append("## SECTION 10 — 🗣️ PLAIN ENGLISH SUMMARY")
    for key, value in summary.items():
        if isinstance(value, dict):
            lines.append(f"- {key}:")
            for subk, subv in value.items():
                lines.append(f"  - {subk}: {subv}")
        else:
            lines.append(f"- {key}: {value}")
    lines.append("")

    lines.append("## SECTION 11 — 📎 ATTACHMENTS, EXHIBITS & REFERENCES")
    if attachments and attachments[0].get("reference") != "NOT FOUND IN DOCUMENT":
        rows = [[a.get("reference", ""), a.get("what_it_is", ""), a.get("attached", "")] for a in attachments]
        lines.append(_markdown_table(["Reference", "What It Is", "Attached? (Yes/No)"], rows))
        for a in attachments:
            lines.append(f"- Locator: {a.get('locator', '')}")
    else:
        lines.append("NOT FOUND IN DOCUMENT")

    lines.append("")
    lines.append("## ⚠️ DISCLAIMER")
    lines.append("This analysis is AI-generated and for informational purposes only. It does not constitute legal advice. Consult a qualified attorney for legal decisions.")
    if ex.extraction_issues:
        lines.append("")
        lines.append("## EXTRACTION ISSUES")
        for issue in ex.extraction_issues:
            lines.append(f"- {issue}")
    return "\n".join(lines).strip()


def build_legal_document_report(
    file_path: str,
    filename: str,
    llm_text_fn: Callable[[str], str] | None = None,
    export_pdf: bool = False,
    export_docx: bool = False,
) -> dict[str, Any]:
    ex = extract_and_prepare(file_path, filename)
    clauses = _extract_clauses_exhaustive(ex)
    alerts = _build_priority_alert_summary(_extract_critical_alerts_exhaustive(ex, clauses))
    report = {
        "document_identity": _extract_document_identity_exhaustive(ex, filename),
        "priority_alerts": alerts,
        "all_parties_involved": _extract_parties_exhaustive(ex),
        "critical_legal_alerts": alerts,
        "all_dates_deadlines": _extract_all_dates_exhaustive(ex),
        "all_financial_details": _extract_financial_details_exhaustive(ex),
        "important_legal_terms_found": _extract_legal_terms_exhaustive(ex, clauses),
        "all_clauses_full_list": clauses,
        "obligations": _extract_obligations_exhaustive(ex, clauses),
        "red_flags_risks": _extract_red_flags_exhaustive(ex, clauses),
        "attachments_exhibits_references": _extract_attachments_exhaustive(ex),
        "raw_extraction": {
            "file_type": ex.file_type,
            "total_pages": ex.total_pages,
            "language": ex.language,
            "word_count": ex.word_count,
            "low_quality_scan": ex.low_quality_scan,
            "ocr_confidence": ex.ocr_confidence,
            "red_flag_terms": ex.red_flag_terms,
            "extraction_issues": ex.extraction_issues,
            "segment_count": len(ex.segments),
        },
    }
    report["plain_english_summary"] = _compose_plain_summary_exhaustive(report)
    report["disclaimer"] = "This analysis is AI-generated and for informational purposes only. It does not constitute legal advice. Consult a qualified attorney for legal decisions."
    markdown = _markdown_from_exhaustive(report, ex)

    exports: dict[str, str] = {}
    if export_pdf:
        pdf_path = os.path.join(tempfile.gettempdir(), f"legal_doc_summary_{os.getpid()}.pdf")
        exports["pdf_path"] = _export_pdf(markdown, pdf_path)
    if export_docx:
        docx_path = os.path.join(tempfile.gettempdir(), f"legal_doc_summary_{os.getpid()}.docx")
        exports["docx_path"] = _export_docx(markdown, docx_path)

    return {
        "status": "ok",
        "filename": filename,
        "file_type": ex.file_type,
        "quality_checks": {
            "word_count": ex.word_count,
            "language": ex.language,
            "low_quality_scan": ex.low_quality_scan,
            "ocr_confidence": ex.ocr_confidence,
            "red_flag_terms": ex.red_flag_terms,
            "extraction_issues": ex.extraction_issues,
            "chunks_used": len(ex.chunks),
            "segment_count": len(ex.segments),
            "requires_human_review": len(ex.red_flag_terms) >= 3,
            "is_large_document": ex.total_pages > 50,
        },
        "extracted_text": {
            "total_pages": ex.total_pages,
            "cleaned_text": ex.cleaned_text,
        },
        "report_json": report,
        "report_markdown": markdown,
        "exports": exports,
    }
