"""Dynamic template processing and DOCX generation for NyayaDraft."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Mapping

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "templates"
GENERATED_DIR = BASE_DIR / "generated"
PLACEHOLDER_REGEX = r"\{\{(.*?)\}\}"
PLACEHOLDER_PATTERN = re.compile(PLACEHOLDER_REGEX)


def list_templates(template_dir: str | Path = TEMPLATE_DIR) -> list[Path]:
    """Return all .txt templates available to the Streamlit dropdown."""
    folder = Path(template_dir)
    folder.mkdir(parents=True, exist_ok=True)
    return sorted(path for path in folder.glob("*.txt") if path.is_file())


def load_template(template_path: str | Path) -> str:
    """Load a selected .txt template."""
    path = Path(template_path)
    if not path.exists():
        raise FileNotFoundError(f"Template not found: {path}")
    if path.suffix.lower() != ".txt":
        raise ValueError(f"Expected a .txt template, got: {path.name}")
    return path.read_text(encoding="utf-8")


def normalize_placeholder_name(name: str) -> str:
    """Normalize a raw placeholder key captured from {{...}}."""
    return re.sub(r"\s+", " ", str(name or "")).strip()


def extract_placeholders(template_text: str) -> list[str]:
    """Extract a unique first-seen list of placeholders using {{...}} regex."""
    placeholders: list[str] = []
    seen: set[str] = set()

    for match in PLACEHOLDER_PATTERN.finditer(template_text):
        placeholder = normalize_placeholder_name(match.group(1))
        if placeholder and placeholder not in seen:
            placeholders.append(placeholder)
            seen.add(placeholder)

    return placeholders


def placeholder_label(name: str) -> str:
    """Convert placeholder_name into a friendly form label."""
    label = normalize_placeholder_name(name).replace("_", " ")
    return label[:1].upper() + label[1:]


def replace_placeholders(
    template_text: str,
    values: Mapping[str, str],
    keep_unfilled: bool = True,
) -> str:
    """Replace {{placeholder}} tags with user-provided values."""
    normalized_values = {
        normalize_placeholder_name(key): str(value or "").strip()
        for key, value in values.items()
    }

    def replace_match(match: re.Match[str]) -> str:
        placeholder = normalize_placeholder_name(match.group(1))
        value = normalized_values.get(placeholder, "")
        if value:
            return value
        return match.group(0) if keep_unfilled else ""

    return PLACEHOLDER_PATTERN.sub(replace_match, template_text)


def missing_placeholders(template_text: str, values: Mapping[str, str]) -> list[str]:
    """Return placeholders that do not have non-empty input values."""
    normalized_values = {
        normalize_placeholder_name(key): str(value or "").strip()
        for key, value in values.items()
    }
    return [
        placeholder
        for placeholder in extract_placeholders(template_text)
        if not normalized_values.get(placeholder)
    ]


def is_all_caps_heading(line: str) -> bool:
    """Return True when a line should be formatted as a Word heading."""
    stripped = line.strip()
    if len(stripped) < 3 or len(stripped) > 140:
        return False

    letters = re.sub(r"[^A-Za-z]", "", stripped)
    if len(letters) < 3:
        return False

    return letters == letters.upper()


def configure_document(document: Document) -> None:
    """Apply simple legal-document page and font defaults."""
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)


def add_body_paragraph(document: Document, line: str) -> None:
    """Add one normal legal-text paragraph."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(line)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def convert_text_to_docx(document_text: str, title: str | None = None) -> BytesIO:
    """Convert final filled text into a BytesIO .docx file."""
    document = Document()
    configure_document(document)

    if title:
        heading = document.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for raw_line in document_text.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
            continue

        if is_all_caps_heading(line):
            heading = document.add_heading(line, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            add_body_paragraph(document, line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def text_to_docx_bytes(document_text: str, title: str | None = None) -> BytesIO:
    """Backward-compatible alias for convert_text_to_docx()."""
    return convert_text_to_docx(document_text=document_text, title=title)


def render_template_to_docx_bytes(
    template_text: str,
    values: Mapping[str, str],
    title: str | None = None,
    keep_unfilled: bool = True,
) -> BytesIO:
    """Fill template text and return the generated .docx as BytesIO."""
    final_text = replace_placeholders(
        template_text=template_text,
        values=values,
        keep_unfilled=keep_unfilled,
    )
    return convert_text_to_docx(document_text=final_text, title=title)


def save_docx(
    document_text: str,
    output_path: str | Path,
    title: str | None = None,
) -> Path:
    """Save final legal text as a .docx file."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    buffer = convert_text_to_docx(document_text=document_text, title=title)
    path.write_bytes(buffer.getvalue())
    return path


def render_template_file_to_docx(
    template_path: str | Path,
    values: Mapping[str, str],
    output_dir: str | Path = GENERATED_DIR,
    keep_unfilled: bool = True,
) -> Path:
    """Fill a template file and save it under generated/ as .docx."""
    template = Path(template_path)
    template_text = load_template(template)
    final_text = replace_placeholders(
        template_text=template_text,
        values=values,
        keep_unfilled=keep_unfilled,
    )
    output_name = f"{template.stem}_draft.docx"
    title = template.stem.replace("_", " ").title()
    return save_docx(final_text, Path(output_dir) / output_name, title=title)
